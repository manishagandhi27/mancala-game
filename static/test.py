import os
import json
import re
from docx import Document
from PIL import Image
import pytesseract
from langchain.document_loaders import UnstructuredWordDocumentLoader

# ----------- Step 1: Process Text, Lists, and Headings -----------

def parse_word_doc(file_path):
    """Extracts headings, paragraphs, and lists from a Word document."""
    doc = Document(file_path)
    extracted_data = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        if para.style.name.startswith("Heading"):
            extracted_data.append({"type": "heading", "content": text})
        elif para.style.name.startswith("List") or re.match(r"^\d+\.", text) or text.startswith("- "):
            extracted_data.append({"type": "list_item", "content": text})
        else:
            extracted_data.append({"type": "paragraph", "content": text})

    return extracted_data

# ----------- Step 2: Extract Tables -----------

def extract_tables(file_path):
    """Extracts tables from a Word document and returns structured table data."""
    doc = Document(file_path)
    tables = []

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append({"type": "table", "content": table_data})

    return tables

# ----------- Step 3: Extract Images & Perform OCR -----------

def extract_images_from_docx(doc_path, output_folder="extracted_images"):
    """Extracts images from a Word document and applies OCR for text extraction."""
    doc = Document(doc_path)
    os.makedirs(output_folder, exist_ok=True)
    
    image_data = []
    for i, shape in enumerate(doc.inline_shapes):
        try:
            img_path = os.path.join(output_folder, f"image_{i}.png")
            with open(img_path, "wb") as f:
                f.write(shape._inline.graphic.graphicData.pic.blipFill.blip.blob)
            
            # Apply OCR to extract text from the image
            img_text = pytesseract.image_to_string(Image.open(img_path)).strip()
            image_data.append({"type": "image", "image_path": img_path, "extracted_text": img_text})
        except Exception as e:
            print(f"Error extracting image {i}: {e}")

    return image_data

# ----------- Step 4: Alternative - Use LangChain Unstructured Loader -----------

def load_using_langchain(file_path):
    """Uses UnstructuredWordDocumentLoader to extract text content."""
    loader = UnstructuredWordDocumentLoader(file_path)
    docs = loader.load()

    structured_data = []
    
    for doc in docs:
        lines = doc.page_content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line.endswith(":"):  # Heading-like text
                structured_data.append({"type": "heading", "content": line})
            elif line.startswith("- ") or re.match(r"^\d+\.", line):  # List detection
                structured_data.append({"type": "list_item", "content": line})
            else:
                structured_data.append({"type": "paragraph", "content": line})

    return structured_data

# ----------- Step 5: Full Document Processing -----------

def process_document(file_path):
    """Combines all extraction functions for a fully processed document."""
    parsed_text = parse_word_doc(file_path)
    tables = extract_tables(file_path)
    images = extract_images_from_docx(file_path)
    
    # Merge results
    final_output = parsed_text + tables + images
    
    return final_output

# ----------- Execute the Code -----------

if __name__ == "__main__":
    file_path = "sample.docx"  # Change to your actual document path
    structured_content = process_document(file_path)
    
    # Pretty print structured output
    print(json.dumps(structured_content, indent=2))


def format_parsed_content(parsed_data):
    """Formats parsed structured content into a readable text format for splitting."""
    formatted_text = []
    
    current_heading = None
    buffer = []
    
    for item in parsed_data:
        if item["type"] == "heading":
            # Store previous section before starting new heading
            if buffer:
                formatted_text.append("\n".join(buffer))
                buffer = []
            current_heading = item["content"]
            buffer.append(f"## {current_heading}")  # Markdown-style heading
            
        elif item["type"] == "list_item":
            buffer.append(f"- {item['content']}")  # Preserve list structure
            
        elif item["type"] == "paragraph":
            buffer.append(item["content"])  # Keep paragraphs grouped
            
    # Append last buffered content
    if buffer:
        formatted_text.append("\n".join(buffer))
    
    return formatted_text




import uuid
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from sentence_transformers import SentenceTransformer

class ChunkProcessor:
    
    def __init__(self, model_name="all-MiniLM-L6-v2"):
        """Initialize the embedding model"""
        self.embedding_model = SentenceTransformer(model_name)

    def generate_chunks(self, documents, chunk_size=500, chunk_overlap=50):
        """
        Splits documents into chunks, generates embeddings, and returns structured chunk data.
        
        Args:
        - documents: List of LangChain Document objects.
        - chunk_size: Max characters per chunk.
        - chunk_overlap: Overlap between chunks to maintain context.
        
        Returns:
        - List of dictionaries containing file_id, chunk_id, chunk_text, and chunk_embedding.
        """
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len
        )
        
        chunked_data = []
        
        for doc in documents:
            file_id = str(uuid.uuid4())  # Unique ID for each document
            chunks = text_splitter.split_text(doc.page_content)
            
            for i, chunk_text in enumerate(chunks):
                chunk_id = f"{file_id}-{i}"  # Unique chunk ID
                chunk_embedding = self.embedding_model.encode(chunk_text).tolist()
                
                chunked_data.append({
                    "file_id": file_id,
                    "chunk_id": chunk_id,
                    "chunk_text": chunk_text,
                    "chunk_embedding": chunk_embedding,
                    "metadata": doc.metadata  # Preserve metadata (e.g., source, page number)
                })
        
        return chunked_data


def process_and_chunk_document(file_path):
    """
    Processes a Word document, converts to LangChain Document format, and chunks it.
    """
    # Step 1: Process Word Document
    parsed_content = DocumentProcessor.process_document(file_path)  # Extracts text, tables, images
    
    # Step 2: Convert to LangChain Document Format
    langchain_docs = convert_to_langchain_documents(parsed_content, file_path)
    
    # Step 3: Chunk the document and generate embeddings
    chunk_processor = ChunkProcessor()
    chunked_docs = chunk_processor.generate_chunks(langchain_docs)
    
    return chunked_docs


file_path = "sample.docx"
chunked_output = process_and_chunk_document(file_path)

# Print sample output
for chunk in chunked_output[:3]:  # Display first 3 chunks
    print(f"File ID: {chunk['file_id']}")
    print(f"Chunk ID: {chunk['chunk_id']}")
    print(f"Chunk Text: {chunk['chunk_text'][:200]}...")  # Show first 200 characters
    print(f"Embedding (First 5 Values): {chunk['chunk_embedding'][:5]}")
    print(f"Metadata: {chunk['metadata']}")
    print("-" * 100)
